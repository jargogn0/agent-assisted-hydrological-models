#!/usr/bin/env python3
"""Build the Supplement for the agent paper: Table S1 (signature definitions,
matching the implementation in paper4_pipeline/src/paper4/metrics.py) and
Figure S1 (spatial distribution of confirmed performance). Writes docx and PDF."""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
FIGS = ROOT / "paper4_pipeline/outputs/hess_100train_50test_final/figures"
FONT = "Times New Roman"

# name, definition, units — worded to match metrics.py exactly
SIGS = [
    ("Mean runoff", "Mean of daily runoff over the evaluation period.", "mm d⁻¹"),
    ("Runoff ratio", "Total runoff divided by total precipitation over the evaluation period.", "–"),
    ("Flow-duration-curve slope",
     "Slope of the log-transformed flow-duration curve between the flows exceeded 33 % and 66 % "
     "of the time: (ln Q33 − ln Q66) / (0.66 − 0.33).", "–"),
    ("Baseflow index",
     "Ratio of baseflow to total flow, with baseflow separated by a three-pass Lyne–Hollick "
     "digital filter (filter parameter 0.925).", "–"),
    ("Half-flow date",
     "Mean day of the water year by which half of the annual runoff volume has passed, "
     "averaged over water years with at least 180 valid days.", "day of year"),
    ("Q95 (high flow)", "95th percentile of the daily runoff distribution.", "mm d⁻¹"),
    ("Q05 (low flow)", "5th percentile of the daily runoff distribution.", "mm d⁻¹"),
    ("High-flow duration",
     "Mean length of consecutive-day spells with runoff at or above the 95th percentile.", "d"),
    ("Low-flow duration",
     "Mean length of consecutive-day spells with runoff at or below the 5th percentile.", "d"),
]


def para(doc, text, size=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         after=8, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    return p


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(2.5))

    para(doc, "Supplement to:", size=11, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "Agent-assisted development of machine-learning hydrological models: "
              "testing a reproducible and controlled workflow", size=14, bold=True,
         after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "Doudou Ba and Jakub Langhammer", size=11, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "Department of Physical Geography and Geoecology, Faculty of Science, "
              "Charles University, Prague, Czech Republic", size=10, italic=True, after=18,
         align=WD_ALIGN_PARAGRAPH.LEFT)

    para(doc, "Table S1. Hydrological signatures retained in the evaluation scorecard: "
              "definitions, computation, and units. All signatures are computed per catchment "
              "and evaluation period from daily runoff, identically for observed and simulated "
              "series; signature errors in the main text are absolute differences between the "
              "simulated and observed values. Definitions correspond exactly to the archived "
              "implementation.", size=10, after=6)

    t = doc.add_table(rows=1 + len(SIGS), cols=3)
    t.style = "Table Grid"
    widths = (Cm(4.2), Cm(9.6), Cm(2.2))
    for j, h in enumerate(("Signature", "Definition", "Units")):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(h)
        r.font.name = FONT; r.font.size = Pt(9.5); r.font.bold = True
    for i, (name, definition, units) in enumerate(SIGS, start=1):
        for j, val in enumerate((name, definition, units)):
            cell = t.rows[i].cells[j]
            cell.paragraphs[0].text = ""
            r = cell.paragraphs[0].add_run(val)
            r.font.name = FONT; r.font.size = Pt(9.5)
    for row in t.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]

    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGS / "figS1_transfer_maps.png"), width=Cm(15.5))
    para(doc, "Figure S1. Catchment-level KGE of the agent-selected configuration of each "
              "model family in the temporal-confirmation period (left column, the 100 "
              "development catchments, 2014–2017) and in the held-out catchments (right "
              "column, the 50 catchments excluded from development, 2014–2017). Warm colours "
              "mark poor performance. Panel medians correspond to the agent entries of "
              "Table 4 in the main text.", size=10, after=0)

    out = ROOT / "AgentPaper_Supplement.docx"
    doc.save(out)
    print("written:", out.name)
    subprocess.run(["/Applications/LibreOffice.app/Contents/MacOS/soffice", "--headless",
                    "--norestore", "--convert-to", "pdf", "--outdir", str(ROOT), str(out)],
                   capture_output=True)
    print("written: AgentPaper_Supplement.pdf")


if __name__ == "__main__":
    main()
